"""
Document Parser Module

Provides layout-aware parsing for PDF, DOCX, DOC, and XLSX documents.
Supports parent-child chunking strategy for improved RAG retrieval:
- Child chunks (small, ~300 tokens): used for precise vector/BM25 search
- Parent chunks (large, ~1500 tokens): used as broad context for the LLM

Layout-aware parsing detects document structure (headers, tables, paragraphs)
to produce semantically coherent chunks aligned to document sections.
"""

import re
import sys
import os
import base64
import logging
from pathlib import Path
from typing import List, Dict, Any
from collections import Counter

import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
from openai import OpenAI

from app.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Vision-based Image Description
# ---------------------------------------------------------------------------

def _describe_image_with_vision(image_bytes: bytes, page_num: int) -> str:
    """
    Use OpenAI Vision API to describe an image extracted from a PDF.

    Args:
        image_bytes: Raw image bytes (PNG/JPEG)
        page_num: Page number for logging

    Returns:
        Textual description of the image, or empty string on failure
    """
    try:
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        client = OpenAI(api_key=settings.openai_api_key)
        response = client.chat.completions.create(
            model=settings.vision_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": settings.vision_prompt,
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{b64}",
                                "detail": "high",
                            },
                        },
                    ],
                }
            ],
            max_completion_tokens=500,
        )
        description = response.choices[0].message.content.strip()
        logger.info(
            "Vision described image on page %d: %.80s...", page_num, description
        )
        return description
    except Exception as e:
        logger.warning("Vision API failed for image on page %d: %s", page_num, e)
        return ""


# ---------------------------------------------------------------------------
# Layout-Aware PDF Parsing
# ---------------------------------------------------------------------------

def _detect_body_font_size(doc: fitz.Document) -> float:
    """
    Scan all pages to determine the most common font size (body text).

    Args:
        doc: PyMuPDF document

    Returns:
        Most common font size (float)
    """
    font_sizes: list[float] = []

    for page in doc:
        try:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        except Exception:
            continue
        for block in blocks:
            if block.get("type") != 0:  # text blocks only
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        font_sizes.append(round(span["size"], 1))

    if not font_sizes:
        return 12.0  # reasonable default

    return Counter(font_sizes).most_common(1)[0][0]


def _format_table_as_markdown(table_data: List[List]) -> str:
    """
    Convert a table (list of rows) to a markdown-style table string.

    Args:
        table_data: List of rows, each row is a list of cell values

    Returns:
        Markdown table string
    """
    if not table_data:
        return ""

    # Clean cell values
    rows = []
    for row in table_data:
        cleaned = []
        for cell in row:
            val = str(cell).strip() if cell is not None else ""
            val = val.replace("|", "\\|")
            cleaned.append(val)
        rows.append(cleaned)

    if not rows:
        return ""

    # Normalize column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    # Build markdown
    lines = []
    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _classify_block_heading_level(
    avg_font_size: float,
    body_size: float,
    is_bold: bool,
    header_font_ratio: float
) -> int:
    """
    Classify a text block as heading level 0 (body), 1, 2, or 3.

    Args:
        avg_font_size: Average font size of the block
        body_size: Body text font size
        is_bold: Whether the block contains bold text
        header_font_ratio: Minimum ratio to classify as heading

    Returns:
        Heading level (0 = body, 1 = H1, 2 = H2, 3 = H3)
    """
    ratio = avg_font_size / body_size if body_size > 0 else 1.0

    if ratio >= 1.5:
        return 1  # H1 - main title
    elif ratio >= header_font_ratio:
        return 2  # H2 - section header
    elif is_bold and ratio > 1.05:
        return 3  # H3 - subsection (bold, slightly larger)
    else:
        return 0  # body text


def parse_pdf_layout_aware(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse PDF with layout-aware structure detection.

    Detects headers by font size, extracts tables as markdown,
    and preserves document hierarchy in the output text.

    Args:
        file_path: Path to the PDF file

    Returns:
        List of page dicts with structured text and metadata
    """
    pages_data = []

    try:
        doc = fitz.open(file_path)
        body_size = _detect_body_font_size(doc)
        header_ratio = settings.header_font_ratio

        for page_num in range(len(doc)):
            page = doc[page_num]
            structured_parts: list[str] = []

            # --- Extract tables first so we can skip their text blocks ---
            table_rects: list[fitz.Rect] = []
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        table_rects.append(fitz.Rect(table.bbox))
                        table_data = table.extract()
                        if table_data:
                            md_table = _format_table_as_markdown(table_data)
                            if md_table:
                                structured_parts.append(md_table)
            except Exception:
                # find_tables may fail on some pages; fall back to text only
                pass

            # --- Process text blocks ---
            try:
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            except Exception:
                # Fallback: plain text extraction
                text = page.get_text()
                if text.strip():
                    pages_data.append({
                        "text": text,
                        "page": page_num + 1,
                        "source": Path(file_path).name
                    })
                continue

            for block in blocks:
                if block.get("type") != 0:
                    # Image block – try to describe with Vision API
                    if block.get("type") == 1:
                        try:
                            xref = block.get("image", 0) or block.get("xref", 0)
                            if xref:
                                img_info = doc.extract_image(xref)
                                if img_info and img_info.get("image"):
                                    img_bytes = img_info["image"]
                                    # Skip very small images (icons, bullets, etc.)
                                    w = block.get("width", 0) or (block["bbox"][2] - block["bbox"][0])
                                    h = block.get("height", 0) or (block["bbox"][3] - block["bbox"][1])
                                    if w > 50 and h > 50:
                                        desc = _describe_image_with_vision(img_bytes, page_num + 1)
                                        if desc:
                                            structured_parts.append(
                                                f"[Immagine/Grafico: {desc}]"
                                            )
                        except Exception as img_err:
                            logger.debug("Could not extract image block: %s", img_err)
                    continue  # skip non-text blocks

                # Check if block overlaps with a table area
                block_rect = fitz.Rect(block["bbox"])
                in_table = any(block_rect.intersects(tr) for tr in table_rects)
                if in_table:
                    continue

                # Collect text and font info from spans
                block_text_parts: list[str] = []
                block_font_sizes: list[float] = []
                block_bold_count = 0
                total_spans = 0

                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        span_text = span.get("text", "")
                        line_text += span_text
                        if span_text.strip():
                            block_font_sizes.append(span["size"])
                            total_spans += 1
                            # PyMuPDF flags: bit 4 = bold (value 16)
                            if span.get("flags", 0) & 16:
                                block_bold_count += 1
                    block_text_parts.append(line_text)

                block_text = "\n".join(block_text_parts).strip()
                if not block_text:
                    continue

                # Classify heading level
                if block_font_sizes:
                    avg_size = sum(block_font_sizes) / len(block_font_sizes)
                    is_bold = block_bold_count > total_spans * 0.5
                    level = _classify_block_heading_level(avg_size, body_size, is_bold, header_ratio)
                else:
                    level = 0

                if level == 1:
                    structured_parts.append(f"# {block_text}")
                elif level == 2:
                    structured_parts.append(f"## {block_text}")
                elif level == 3:
                    structured_parts.append(f"### {block_text}")
                else:
                    structured_parts.append(block_text)

            page_text = "\n\n".join(structured_parts)
            if page_text.strip():
                pages_data.append({
                    "text": page_text,
                    "page": page_num + 1,
                    "source": Path(file_path).name
                })

        doc.close()

    except Exception as e:
        raise ValueError(f"Error parsing PDF (layout-aware): {str(e)}")

    return pages_data


# ---------------------------------------------------------------------------
# Standard (fallback) Parsers
# ---------------------------------------------------------------------------

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse PDF document and extract text with page information (simple mode).

    Args:
        file_path: Path to the PDF file

    Returns:
        List of dictionaries containing text and metadata for each page
    """
    pages_data = []

    try:
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            if text.strip():
                pages_data.append({
                    "text": text,
                    "page": page_num + 1,
                    "source": Path(file_path).name
                })

        doc.close()

    except Exception as e:
        raise ValueError(f"Error parsing PDF: {str(e)}")

    return pages_data


def parse_docx_layout_aware(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse Word document (.docx) with layout-aware structure detection.
    Uses paragraph styles to detect headings.

    Args:
        file_path: Path to the Word file

    Returns:
        List of dictionaries containing structured text and metadata
    """
    try:
        doc = Document(file_path)

        structured_parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()

            if "heading 1" in style_name or "titolo 1" in style_name:
                structured_parts.append(f"# {text}")
            elif "heading 2" in style_name or "titolo 2" in style_name:
                structured_parts.append(f"## {text}")
            elif "heading 3" in style_name or "titolo 3" in style_name:
                structured_parts.append(f"### {text}")
            elif "heading" in style_name or "titolo" in style_name:
                # Generic heading
                structured_parts.append(f"## {text}")
            elif "title" in style_name:
                structured_parts.append(f"# {text}")
            else:
                structured_parts.append(text)

        # Also extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                md_table = _format_table_as_markdown(table_data)
                if md_table:
                    structured_parts.append(md_table)

        full_text = "\n\n".join(structured_parts)

        return [{
            "text": full_text,
            "page": 1,
            "source": Path(file_path).name
        }]

    except Exception as e:
        raise ValueError(f"Error parsing DOCX (layout-aware): {str(e)}")


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse Word document (.docx) and extract text (simple mode).

    Args:
        file_path: Path to the Word file

    Returns:
        List of dictionaries containing text and metadata
    """
    try:
        doc = Document(file_path)

        full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        return [{
            "text": full_text,
            "page": 1,
            "source": Path(file_path).name
        }]

    except Exception as e:
        raise ValueError(f"Error parsing DOCX: {str(e)}")


def parse_doc(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse old Word document (.doc) format using pywin32 on Windows.

    Args:
        file_path: Path to the .doc file

    Returns:
        List of dictionaries containing text and metadata
    """
    if sys.platform != 'win32':
        raise ValueError("Parsing .doc files is only supported on Windows. Please convert to .docx format.")

    try:
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()

        try:
            abs_path = os.path.abspath(file_path)

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                doc = word.Documents.Open(abs_path, ReadOnly=True)
                full_text = doc.Content.Text
                full_text = full_text.replace('\r', '\n')
                doc.Close(False)

                return [{
                    "text": full_text,
                    "page": 1,
                    "source": Path(file_path).name
                }]
            finally:
                word.Quit()
        finally:
            pythoncom.CoUninitialize()

    except ImportError:
        raise ValueError("pywin32 is required to parse .doc files. Install with: pip install pywin32")
    except Exception as e:
        raise ValueError(f"Error parsing DOC: {str(e)}")


def parse_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse Excel document (.xlsx) and extract text from all sheets.

    Args:
        file_path: Path to the Excel file

    Returns:
        List of dictionaries containing text and metadata for each sheet
    """
    try:
        wb = load_workbook(file_path, data_only=True)
        sheets_data = []

        for sheet_idx, sheet_name in enumerate(wb.sheetnames, 1):
            sheet = wb[sheet_name]
            rows_text = []

            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    rows_text.append(" | ".join(row_values))

            if rows_text:
                sheet_text = f"--- Foglio: {sheet_name} ---\n" + "\n".join(rows_text)
                sheets_data.append({
                    "text": sheet_text,
                    "page": sheet_idx,
                    "source": Path(file_path).name
                })

        wb.close()
        return sheets_data

    except Exception as e:
        raise ValueError(f"Error parsing XLSX: {str(e)}")


# ---------------------------------------------------------------------------
# Chunking: Standard (legacy) and Parent-Child
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
    """
    Split text into chunks with overlap, trying to maintain semantic boundaries.
    Legacy method – used when parent-child chunking is disabled.

    Args:
        text: Text to chunk
        chunk_size: Maximum characters per chunk (default from settings)
        chunk_overlap: Characters of overlap between chunks (default from settings)

    Returns:
        List of text chunks
    """
    if chunk_size is None:
        chunk_size = settings.chunk_size
    if chunk_overlap is None:
        chunk_overlap = settings.chunk_overlap

    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks: list[str] = []
    current_chunk: list[str] = []
    current_length = 0

    for sentence in sentences:
        sentence_length = len(sentence)

        if current_length + sentence_length > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))

            # Build overlap from the end of the previous chunk
            overlap_sentences: list[str] = []
            overlap_length = 0
            for sent in reversed(current_chunk):
                if overlap_length + len(sent) <= chunk_overlap:
                    overlap_sentences.insert(0, sent)
                    overlap_length += len(sent)
                else:
                    break

            current_chunk = overlap_sentences
            current_length = overlap_length

        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


def _split_into_sections(text: str) -> List[Dict[str, Any]]:
    """
    Split structured text into sections based on markdown-style headers.
    Each section contains a heading path and its body content.

    Args:
        text: Structured text with markdown headers

    Returns:
        List of sections: [{"heading_path": "...", "body": "...", "level": int}]
    """
    lines = text.split("\n")
    sections: list[Dict[str, Any]] = []

    # Track current heading hierarchy
    headings: Dict[int, str] = {}
    current_body_lines: list[str] = []
    current_level = 0

    def flush_section():
        body = "\n".join(current_body_lines).strip()
        if body or headings:
            path_parts = [headings[k] for k in sorted(headings.keys()) if headings.get(k)]
            heading_path = " > ".join(path_parts) if path_parts else ""
            sections.append({
                "heading_path": heading_path,
                "body": body,
                "level": current_level
            })

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_body_lines.append("")
            continue

        # Detect heading level
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            # Flush previous section
            flush_section()
            current_body_lines = []

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            current_level = level

            # Update heading hierarchy
            headings[level] = heading_text
            # Clear deeper headings
            for deeper in list(headings.keys()):
                if deeper > level:
                    del headings[deeper]
        else:
            current_body_lines.append(line)

    # Flush last section
    flush_section()

    return sections


def create_parent_child_chunks(
    text: str,
    parent_size: int = None,
    child_size: int = None,
    child_overlap: int = None,
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Create parent and child chunks from structured text.

    Parent chunks are large (~1500 tokens) and section-aligned.
    Child chunks are small (~300 tokens) derived from parents for precise search.

    Args:
        text: Structured text (with markdown headers from layout-aware parsing)
        parent_size: Max characters for parent chunks
        child_size: Max characters for child chunks
        child_overlap: Character overlap between child chunks

    Returns:
        Dict with 'parent_chunks' and 'child_chunks' lists
    """
    if parent_size is None:
        parent_size = settings.parent_chunk_size * 4  # ~4 chars per token
    if child_size is None:
        child_size = settings.child_chunk_size * 4
    if child_overlap is None:
        child_overlap = settings.child_chunk_overlap * 4

    sections = _split_into_sections(text)

    parent_chunks: list[Dict[str, Any]] = []
    child_chunks: list[Dict[str, Any]] = []

    # --- Build parent chunks by accumulating sections ---
    current_parent_text = ""
    current_parent_heading = ""
    parent_index = 0

    def flush_parent():
        nonlocal current_parent_text, current_parent_heading, parent_index
        text_stripped = current_parent_text.strip()
        if not text_stripped:
            return

        # Prepend heading context for the parent chunk
        if current_parent_heading:
            full_text = f"[Sezione: {current_parent_heading}]\n{text_stripped}"
        else:
            full_text = text_stripped

        parent_chunks.append({
            "text": full_text,
            "parent_index": parent_index
        })

        # PASS text_stripped here, NOT full_text, to avoid double-heading in the first child!
        _create_children_from_parent(
            text_stripped, parent_index, child_size, child_overlap, child_chunks, current_parent_heading
        )

        parent_index += 1
        current_parent_text = ""

    # --- Iterate over sections and accumulate into parent chunks ---
    for section in sections:
        section_body = section["body"]
        heading = section["heading_path"]

        # Skip empty sections, but track heading changes
        if not section_body.strip():
            if heading:
                current_parent_heading = heading
            continue

        # If adding this section would exceed parent_size, flush current parent
        addition_len = len(section_body) + (2 if current_parent_text else 0)
        if current_parent_text and len(current_parent_text) + addition_len > parent_size:
            flush_parent()

        # Update heading from section
        if heading:
            current_parent_heading = heading

        # Accumulate section body into current parent
        if current_parent_text:
            current_parent_text += "\n\n" + section_body
        else:
            current_parent_text = section_body

    # Flush remaining
    flush_parent()

    return {
        "parent_chunks": parent_chunks,
        "child_chunks": child_chunks
    }


def _create_children_from_parent(
    parent_text: str,
    parent_index: int,
    child_size: int,
    child_overlap: int,
    child_chunks: list,
    parent_heading: str = "",
):
    """
    Split a parent chunk into overlapping child chunks.

    Args:
        parent_text: Text of the parent chunk
        parent_index: Index of the parent chunk
        child_size: Max characters per child chunk
        child_overlap: Character overlap between children
        child_chunks: List to append children to (mutated in place)
        parent_heading: Optional heading to prepend to each child for
                        better vector-search context (avoids context starvation)
    """
    sentences = re.split(r'(?<=[.!?])\s+', parent_text)

    current_child: list[str] = []
    current_length = 0

    for sentence in sentences:
        sentence_len = len(sentence)

        if current_length + sentence_len > child_size and current_child:
            child_text = " ".join(current_child)
            if parent_heading:
                child_text = f"[Sezione: {parent_heading}]\n" + child_text
            child_chunks.append({
                "text": child_text,
                "parent_index": parent_index,
                "child_index": len(child_chunks)
            })

            # Overlap: keep tail sentences
            overlap_sents: list[str] = []
            overlap_len = 0
            for s in reversed(current_child):
                if overlap_len + len(s) <= child_overlap:
                    overlap_sents.insert(0, s)
                    overlap_len += len(s)
                else:
                    break

            current_child = overlap_sents
            current_length = overlap_len

        current_child.append(sentence)
        current_length += sentence_len

    # Last child
    if current_child:
        child_text = " ".join(current_child)
        if parent_heading:
            child_text = f"[Sezione: {parent_heading}]\n" + child_text
        child_chunks.append({
            "text": child_text,
            "parent_index": parent_index,
            "child_index": len(child_chunks)
        })


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def parse_document(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Parse document and return chunked text with metadata.

    When parent-child chunking is enabled, returns both parent and child chunks.
    Otherwise returns a flat list of chunks (legacy mode).

    Args:
        file_path: Path to the document file
        file_type: Type of file ('pdf', 'docx', 'doc', 'xlsx')

    Returns:
        Dictionary with:
        - "parent_chunks": list of parent chunks with metadata
        - "child_chunks": list of child chunks with metadata (these get embedded)
        - "mode": "parent_child" or "legacy"
    """
    layout_aware = settings.enable_layout_aware_parsing
    parent_child = settings.enable_parent_child

    # --- Parse document based on type ---
    ft = file_type.lower()
    if ft == 'pdf':
        pages_data = parse_pdf_layout_aware(file_path) if layout_aware else parse_pdf(file_path)
    elif ft == 'docx':
        pages_data = parse_docx_layout_aware(file_path) if layout_aware else parse_docx(file_path)
    elif ft == 'doc':
        pages_data = parse_doc(file_path)
    elif ft == 'xlsx':
        pages_data = parse_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    if parent_child:
        return _build_parent_child_chunks(pages_data)
    else:
        return _build_legacy_chunks(pages_data)


def _build_parent_child_chunks(pages_data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build parent and child chunks from parsed pages.

    Args:
        pages_data: List of page dictionaries with text and metadata

    Returns:
        Dict with parent_chunks, child_chunks, and mode
    """
    all_parent_chunks: list[Dict[str, Any]] = []
    all_child_chunks: list[Dict[str, Any]] = []

    # Process all pages together to create coherent parent chunks
    for page_data in pages_data:
        result = create_parent_child_chunks(page_data["text"])

        page_num = page_data["page"]
        source = page_data["source"]

        # Re-index parents globally
        parent_offset = len(all_parent_chunks)

        for pc in result["parent_chunks"]:
            all_parent_chunks.append({
                "text": pc["text"],
                "metadata": {
                    "source": source,
                    "page": page_num,
                    "parent_chunk_index": parent_offset + pc["parent_index"]
                }
            })

        # Re-index children globally with correct parent references
        for cc in result["child_chunks"]:
            all_child_chunks.append({
                "text": cc["text"],
                "metadata": {
                    "source": source,
                    "page": page_num,
                    "chunk_index": len(all_child_chunks),
                    "parent_chunk_index": parent_offset + cc["parent_index"]
                }
            })

    return {
        "parent_chunks": all_parent_chunks,
        "child_chunks": all_child_chunks,
        "mode": "parent_child"
    }


def _build_legacy_chunks(pages_data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build flat chunks (legacy mode, no parent-child).

    Args:
        pages_data: List of page dictionaries

    Returns:
        Dict with child_chunks only (for backward compatibility) and mode
    """
    all_chunks: list[Dict[str, Any]] = []
    chunk_index = 0

    for page_data in pages_data:
        page_chunks = chunk_text(page_data["text"])

        for chunk in page_chunks:
            all_chunks.append({
                "text": chunk,
                "metadata": {
                    "source": page_data["source"],
                    "page": page_data["page"],
                    "chunk_index": chunk_index
                }
            })
            chunk_index += 1

    return {
        "parent_chunks": [],
        "child_chunks": all_chunks,
        "mode": "legacy"
    }


def get_document_page_count(file_path: str, file_type: str) -> int:
    """
    Get the number of pages in a document.

    Args:
        file_path: Path to the document file
        file_type: Type of file ('pdf', 'docx', 'xlsx')

    Returns:
        Number of pages (or sheets for Excel)
    """
    if file_type.lower() == 'pdf':
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)
            doc.close()
            return page_count
        except Exception:
            return 0
    elif file_type.lower() in ['docx', 'doc']:
        return 1
    elif file_type.lower() == 'xlsx':
        try:
            wb = load_workbook(file_path, read_only=True)
            sheet_count = len(wb.sheetnames)
            wb.close()
            return sheet_count
        except Exception:
            return 0
    else:
        return 0
